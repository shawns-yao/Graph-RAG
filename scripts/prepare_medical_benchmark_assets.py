#!/usr/bin/env python3
"""Clean medical benchmark JSON files and export md/docx/pdf corpus assets."""

from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path
from typing import Any

from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

ROOT = Path(__file__).resolve().parent.parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

BENCH_DIR = ROOT / "test" / "medical_benchmark"
DOC_EXPORT_DIR = BENCH_DIR / "corpus_exports"

QUESTION_FILES = [
    "questions_master.json",
    "questions_doc_001_simple.json",
    "questions_doc_001_relation.json",
    "questions_doc_001_multihop.json",
    "questions_doc_001_global_temporal.json",
    "questions_doc_001_part1.json",
    "ab_graph_advantage.json",
    "ab_lexical_advantage.json",
    "ab_self_correction.json",
    "ab_semantic_advantage.json",
]


def _load_json(path: Path) -> Any:
    return json.loads(path.read_text(encoding="utf-8"))


def _dump_json(path: Path, payload: Any) -> None:
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def _extract_json_segments(text: str) -> list[Any]:
    payloads: list[Any] = []
    idx = 0
    length = len(text)
    while idx < length:
        while idx < length and text[idx] not in "[{":
            idx += 1
        if idx >= length:
            break
        start = idx
        opening = text[idx]
        closing = "]" if opening == "[" else "}"
        depth = 0
        in_string = False
        escape = False
        end = -1
        for pos in range(idx, length):
            ch = text[pos]
            if in_string:
                if escape:
                    escape = False
                elif ch == "\\":
                    escape = True
                elif ch == '"':
                    in_string = False
                continue
            if ch == '"':
                in_string = True
                continue
            if ch == opening:
                depth += 1
            elif ch == closing:
                depth -= 1
                if depth == 0:
                    end = pos
                    break
        if end < 0:
            break
        snippet = text[start : end + 1]
        try:
            payloads.append(json.loads(snippet))
        except json.JSONDecodeError:
            pass
        idx = end + 1
    return payloads


def _flatten_question_payload(payloads: list[Any]) -> list[dict[str, Any]]:
    rows: list[dict[str, Any]] = []
    for payload in payloads:
        if isinstance(payload, list):
            rows.extend(item for item in payload if isinstance(item, dict))
        elif isinstance(payload, dict):
            if "questions" in payload and isinstance(payload["questions"], list):
                rows.extend(item for item in payload["questions"] if isinstance(item, dict))
            else:
                rows.append(payload)
    return rows


def _dedupe_rows(rows: list[dict[str, Any]], *, key_fields: tuple[str, ...]) -> list[dict[str, Any]]:
    seen: set[tuple[Any, ...]] = set()
    deduped: list[dict[str, Any]] = []
    for row in rows:
        key = tuple(row.get(field) for field in key_fields)
        if key in seen:
            continue
        seen.add(key)
        deduped.append(row)
    return deduped


def clean_question_files() -> list[dict[str, Any]]:
    cleaned_summary: list[dict[str, Any]] = []
    for name in QUESTION_FILES:
        path = BENCH_DIR / name
        if not path.exists():
            continue
        text = path.read_text(encoding="utf-8")
        payloads = _extract_json_segments(text)
        if name == "questions_master.json":
            metadata = {}
            rows = []
            for payload in payloads:
                if isinstance(payload, dict) and "questions" in payload:
                    metadata = payload.get("metadata", {})
                    rows.extend(item for item in payload.get("questions", []) if isinstance(item, dict))
                elif isinstance(payload, dict):
                    rows.append(payload)
            rows = _dedupe_rows(rows, key_fields=("id", "query"))
            if metadata:
                metadata["total_questions"] = len(rows)
                query_types: dict[str, int] = {}
                for row in rows:
                    qtype = str(row.get("query_type") or "unknown")
                    query_types[qtype] = query_types.get(qtype, 0) + 1
                metadata["query_types"] = query_types
            _dump_json(path, {"metadata": metadata, "questions": rows})
            cleaned_summary.append({"file": name, "rows": len(rows)})
            continue

        rows = _flatten_question_payload(payloads)
        rows = _dedupe_rows(rows, key_fields=("question_id", "question"))
        _dump_json(path, rows)
        cleaned_summary.append({"file": name, "rows": len(rows)})
    return cleaned_summary


def _markdown_to_plain_lines(markdown: str) -> list[str]:
    lines: list[str] = []
    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            lines.append("")
            continue
        if line.lstrip().startswith("#"):
            text = line.lstrip("#").strip()
            lines.append(text)
            continue
        if line.lstrip().startswith("- "):
            lines.append(f"• {line.lstrip()[2:].strip()}")
            continue
        lines.append(line)
    return lines


def _render_doc_markdown(doc: dict[str, Any]) -> str:
    title = str(doc.get("title") or doc.get("doc_id") or "Medical Document").strip()
    body = str(doc.get("body_markdown") or "").strip()
    sections: list[str] = [f"# {title}", "", body]

    tables = doc.get("tables") or []
    if tables:
        sections.extend(["", "## Extracted Tables"])
        for table in tables:
            caption = str(table.get("caption") or table.get("table_id") or "Table").strip()
            markdown = str(table.get("markdown") or "").strip()
            sections.extend(["", f"### {caption}", "", markdown])
            for fact in table.get("key_facts") or []:
                sections.append(f"- {fact}")

    entities = doc.get("key_entities") or []
    if entities:
        sections.extend(["", "## Key Entities"])
        for entity in entities:
            aliases = ", ".join(entity.get("aliases") or [])
            line = f"- {entity.get('name', '')} [{entity.get('entity_type', '')}]"
            if aliases:
                line += f" aliases: {aliases}"
            sections.append(line)

    relations = doc.get("key_relations") or []
    if relations:
        sections.extend(["", "## Key Relations"])
        for relation in relations:
            sections.append(
                f"- {relation.get('source', '')} --{relation.get('relation', '')}--> {relation.get('target', '')}"
            )

    return "\n".join(part for part in sections if part is not None).strip() + "\n"


def _write_markdown(path: Path, markdown: str) -> None:
    path.write_text(markdown, encoding="utf-8")


def _write_docx(path: Path, markdown: str) -> None:
    document = Document()
    for line in _markdown_to_plain_lines(markdown):
        stripped = line.strip()
        if not stripped:
            document.add_paragraph("")
            continue
        if stripped.startswith("• "):
            document.add_paragraph(stripped[2:], style="List Bullet")
            continue
        document.add_paragraph(stripped)
    document.save(path)


def _write_pdf(path: Path, markdown: str) -> None:
    styles = getSampleStyleSheet()
    body_style = styles["BodyText"]
    title_style = styles["Heading1"]
    section_style = styles["Heading2"]
    story: list[Any] = []

    for line in _markdown_to_plain_lines(markdown):
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 8))
            continue
        style = body_style
        if line == stripped and len(stripped) <= 32:
            if not story:
                style = title_style
            elif re.match(r"^\d+(\.\d+)*", stripped):
                style = section_style
        text = stripped.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        story.append(Paragraph(text, style))
        story.append(Spacer(1, 4))

    path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(str(path), pagesize=A4, leftMargin=48, rightMargin=48, topMargin=48, bottomMargin=48)
    doc.build(story)


def export_corpus_assets() -> list[dict[str, str]]:
    payload = _load_json(BENCH_DIR / "corpus_documents.json")
    if not isinstance(payload, list):
        raise ValueError("corpus_documents.json must contain a JSON array")

    DOC_EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    for existing in DOC_EXPORT_DIR.glob("*"):
        if existing.is_file():
            existing.unlink()

    exports: list[dict[str, str]] = []
    for doc in payload:
        doc_id = str(doc.get("doc_id") or "medical_doc").strip()
        markdown = _render_doc_markdown(doc)
        md_path = DOC_EXPORT_DIR / f"{doc_id}.md"
        docx_path = DOC_EXPORT_DIR / f"{doc_id}.docx"
        pdf_path = DOC_EXPORT_DIR / f"{doc_id}.pdf"
        _write_markdown(md_path, markdown)
        _write_docx(docx_path, markdown)
        _write_pdf(pdf_path, markdown)
        exports.append(
            {
                "doc_id": doc_id,
                "md": str(md_path),
                "docx": str(docx_path),
                "pdf": str(pdf_path),
            }
        )
    return exports


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--skip-clean", action="store_true", help="Skip question JSON cleanup")
    parser.add_argument("--skip-export", action="store_true", help="Skip corpus export")
    args = parser.parse_args()

    payload: dict[str, Any] = {}
    if not args.skip_clean:
        payload["cleaned_questions"] = clean_question_files()
    if not args.skip_export:
        payload["exported_documents"] = export_corpus_assets()
        payload["export_dir"] = str(DOC_EXPORT_DIR)
    print(json.dumps(payload, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
